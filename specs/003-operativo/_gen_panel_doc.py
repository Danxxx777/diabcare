# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTS = [
    Path(r"D:\6to Software\Construcción de Software\EVF 14.docx"),
    Path(r"D:\6to Software\Construcción de Software\DiabCare-Panel-Workpanel-ELT.docx"),
    Path(r"C:\Users\USER\Downloads\EVF 14.docx"),
]

def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = color

def p(doc, text, size=11, bold=False, after=8):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = para.add_run(text); font(r, size, bold); return para

def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs: font(r, 14 if level==1 else 12, True, RGBColor(0x0F,0x17,0x2A))
    return x

def bullets(doc, items):
    for it in items:
        para = doc.add_paragraph(it, style="List Bullet")
        for r in para.runs: font(r, 11)
        para.paragraph_format.space_after = Pt(2)

def cell(c, text, bold=False, size=9):
    c.text = ""; r = c.paragraphs[0].add_run(text); font(r, size, bold)

def shade(row):
    for c in row.cells:
        tc = c._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), "E2E8F0"); sh.set(qn("w:val"), "clear"); tc.append(sh)

def table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    for j,h_ in enumerate(headers): cell(t.rows[0].cells[j], h_, True, 9)
    shade(t.rows[0])
    for i,row in enumerate(rows,1):
        for j,v in enumerate(row): cell(t.rows[i].cells[j], str(v), False, 9)
    doc.add_paragraph()

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(1.8); s.bottom_margin=Cm(1.8); s.left_margin=Cm(2); s.right_margin=Cm(2)

    p(doc, "Est. Byron Loor Mendoza", 11, True, 2)
    p(doc, "Link de Video: https://youtu.be/3425-TO3UN4", 11, False, 2)
    p(doc, "Repositorio: https://github.com/Danxxx777/diabcare.git", 11, False, 2)
    p(doc, "DiabCare Analytics — Operativo completo · Táctico completo · Estratégico AGG · ELT", 12, True, 10)

    h(doc, "1. Introducción y alcance")
    p(doc, "DiabCare Analytics es una plataforma web hospitalaria de analítica clínica en diabetes. La entrega demuestra tres niveles empresariales en el sistema en ejecución:")
    bullets(doc, [
        "Operativo completo: ejecución diaria con CRUD real (expediente, agenda, atención, laboratorio, farmacia, caja, gobierno).",
        "Táctico completo: Panel con Workpanel, informes simples e informes compuestos, más PDF y vistas BI.",
        "Estratégico AGG: lectura de dirección mediante agregados (totales, promedios, rankings, márgenes) materializados en MinIO (agg_*). AGG = agregar/agregado.",
        "ELT: puente Extract → Load → Transform que alimenta táctica y AGG sin romper el operativo.",
    ])
    p(doc, "El Panel unifica la lectura táctica y AGG; los módulos operativos son donde se escribe y atiende. No se confunde un listado CRUD con un agregado estratégico.")

    h(doc, "2. Arquitectura mínima de referencia")
    bullets(doc, [
        "Frontend: frontend/paginas/ por departamento + estaticos (navegación, i18n, API).",
        "Backend: FastAPI en backend/paquetes/ (auth, clínico, dataset, farmacia, facturación, etc.).",
        "Datos: MinIO/Parquet — stage/ (clínico), negocio/ (operación + agg_*), agregados/ y hechos/dims del DWH.",
        "Orquestación: Apache Airflow embebido (Datos → Orquestador), DAG diabcare_elt.",
        "ML: Random Forest (predicción de riesgo) y métricas en Modelo ML.",
    ])

    h(doc, "3. Nivel operativo (completo)")
    p(doc, "Completo significa: autenticación, autorización por rol, UI + API + persistencia Parquet por módulo del flujo hospitalario. No son pantallas vacías.")
    h(doc, "3.1 Flujo operativo demostrable", 2)
    bullets(doc, [
        "Login → sesión por cookie httpOnly.",
        "Pacientes: crear/buscar expediente (cédula, nombre, foto).",
        "Admisiones: registrar ingreso; Agenda: asignar cita; Mis citas (médico): confirmar/atender.",
        "Registros clínicos / Urgencias / Laboratorio / Comorbilidades: documentar encuentro y órdenes.",
        "Recetas → Farmacia (dispensar, stock) → Facturación/Caja (cobrar, recibo).",
        "Notificaciones y Auditoría: alertas y trazas; Usuarios/Configuración: gobierno.",
    ])
    h(doc, "3.2 Catálogo de módulos operativos", 2)
    table(doc, ["Depto", "Módulo", "Ruta UI", "API típica", "Estado"], [
        ["Seguridad", "Login / sesión", "/", "/api/auth/*", "Implementado"],
        ["Seguridad", "Usuarios", "seguridad/usuarios", "/api/usuarios", "Implementado"],
        ["Seguridad", "Perfil", "seguridad/perfil", "/api/auth/perfil", "Implementado"],
        ["Seguridad", "Notificaciones", "seguridad/notificaciones", "/api/notificaciones", "Implementado"],
        ["Clínico", "Pacientes HCE", "clinico/pacientes", "/api/pacientes", "Implementado"],
        ["Clínico", "Admisiones", "clinico/admisiones", "/api/admisiones", "Implementado"],
        ["Clínico", "Agenda", "clinico/agenda", "/api/citas", "Implementado"],
        ["Clínico", "Mis citas", "clinico/mis_citas", "/api/mis-citas", "Implementado"],
        ["Clínico", "Registros clínicos", "clinico/registros_clinicos", "/api/registros", "Implementado"],
        ["Clínico", "Urgencias", "clinico/urgencias", "/api/urgencias", "Implementado"],
        ["Clínico", "Laboratorio", "clinico/laboratorio", "/api/laboratorio", "Implementado"],
        ["Clínico", "Comorbilidades", "clinico/comorbilidades", "/api/comorbilidades", "Implementado"],
        ["Farmacia", "Mostrador / inventario", "negocio/farmacia", "/api/farmacia", "Implementado"],
        ["Farmacia", "Recetas", "negocio/recetas", "/api/recetas", "Implementado"],
        ["Negocio", "Facturación / caja", "negocio/facturacion", "/api/facturacion", "Implementado"],
        ["Negocio", "RRHH / costeo", "negocio/rrhh", "/api/rrhh", "Implementado"],
        ["Gobierno", "Auditoría", "gobierno/auditoria", "/api/auditoria", "Implementado"],
        ["Gobierno", "Configuración", "gobierno/configuracion", "/api/configuracion", "Implementado"],
    ])
    h(doc, "3.3 Matriz de permisos (extracto)", 2)
    table(doc, ["Capacidad", "Admin", "Médico", "Analista", "Farmacéutico"], [
        ["Usuarios / config / auditoría", "Sí", "No", "No", "No"],
        ["Pacientes / admisiones / agenda", "Sí", "Atención / lectura", "No", "No"],
        ["Mis citas / registros / urgencias / lab", "Sí", "Sí", "No", "No"],
        ["Dataset / pipeline / modelo", "Sí", "No", "Sí", "No"],
        ["Panel / reportes / predicción", "Sí", "Sí", "Sí", "Parcial"],
        ["Farmacia / recetas / caja", "Sí", "Recetas", "No", "Sí"],
    ])
    p(doc, "Criterio de cierre operativo: recorrido paciente → cita → atención → lab/farmacia → cobro deja datos en MinIO y aparece en resúmenes del Panel.")

    h(doc, "4. Nivel táctico (completo)")
    p(doc, "Gestiona el periodo: pulso, listados del día y cruces. Se concentra en Análisis → Panel, Exportar PDF, Estadísticas, Calidad diabetes, Dataset.")
    h(doc, "4.1 Workpanel", 2)
    p(doc, "Tablero del Panel (/api/registros/estadisticas): total de encuentros, con diabetes, sin diabetes, prevalencia. Prioriza hacia un informe simple o compuesto.")
    h(doc, "4.2 Informes simples", 2)
    p(doc, "Criterio: capa transaccional (listado/filtro/detalle). En el Panel: tarjetas-resumen + enlace al módulo.")
    table(doc, ["N.°", "Informe simple", "Endpoint / origen", "Dónde"], [
        ["1", "Agenda / Mis citas (hoy)", "/api/citas/resumen o /api/mis-citas", "Panel + módulo"],
        ["2", "Urgencias", "/api/urgencias/resumen", "Panel + módulo"],
        ["3", "Laboratorio", "/api/laboratorio/resumen", "Panel + módulo"],
        ["4", "Caja / facturas", "/api/facturacion/resumen", "Panel + módulo"],
        ["5", "Padrón pacientes", "/api/pacientes/resumen", "Panel + módulo"],
        ["6", "Admisiones", "/api/admisiones/resumen", "Panel + módulo"],
        ["7", "Farmacia", "/api/farmacia/resumen", "Panel + módulo"],
        ["8", "RRHH", "/api/rrhh/resumen", "Panel + módulo"],
        ["9", "Comorbilidades", "/api/comorbilidades/resumen", "Panel + módulo"],
        ["10", "Factura/recibo imprimible", "Caja / Farmacia", "Panel (atajo)"],
        ["11", "Registros clínicos", "/api/registros", "Módulo"],
        ["12", "Recetas", "módulo recetas", "Módulo"],
        ["13", "Usuarios / sesiones", "/api/usuarios", "Módulo"],
        ["14", "Auditoría", "/api/auditoria", "Módulo"],
        ["15", "Notificaciones", "/api/notificaciones", "Módulo / topbar"],
    ])
    h(doc, "4.3 Informes compuestos", 2)
    p(doc, "Criterio: requieren agregar (GROUP BY / agg_*). API: /api/dataset/informes/complejos.")
    table(doc, ["N.°", "Compuesto en Panel", "Fuente AGG / hechos"], [
        ["1", "Ingresos por día", "Facturas / hechos facturación"],
        ["2", "Top medicamentos dispensados", "agg_medicamentos_top"],
        ["3", "Productividad médica", "agg_productividad_medica"],
        ["4", "Margen farmacia", "agg_margen_farmacia"],
        ["5", "Costo vs facturado por servicio", "agg_costo_servicio"],
        ["6", "Espera urgencias", "agg_tiempos_espera"],
    ])
    p(doc, "Complementos: Exportar PDF (simple/compuesto/completo), Dataset DWH, Estadísticas, Calidad diabetes. Sin datos hospitalarios o sin materialización, compuestos vacíos.")
    h(doc, "4.4 Criterio de cierre táctico", 2)
    bullets(doc, [
        "Workpanel + simples + compuestos en el Panel.",
        "Distinción explícita simple (operación) vs compuesto (agregación).",
        "PDF exportable y drill-down a módulo.",
        "Compuestos trazables a MinIO agg_* / hechos.",
    ])

    h(doc, "5. Nivel estratégico (AGG)")
    p(doc, "Estratégico AGG = dirección con indicadores ya agregados. AGG = agregar → agregado (totales, promedios, rankings, márgenes). No es un listado operativo.")
    h(doc, "5.1 AGG en el Panel (dirección)", 2)
    p(doc, "Bloque Estratégico AGG → /api/dataset/negocio/kpis:")
    table(doc, ["Indicador AGG", "Pregunta de dirección", "Origen"], [
        ["Facturado total", "¿Cuánto ingresó el hospital?", "Facturación / hechos"],
        ["Margen farmacia", "¿La dispensación es rentable?", "agg_margen_farmacia"],
        ["Espera urgencias", "¿Hay saturación de respuesta?", "agg_tiempos_espera"],
        ["Productividad", "¿Cuál es la carga asistencial agregada?", "agg_productividad_medica"],
        ["Lab pendiente", "¿Hay cuello diagnóstico?", "Resumen laboratorio"],
    ])
    h(doc, "5.2 Catálogo AGG en MinIO / DWH", 2)
    table(doc, ["Tabla AGG", "Ubicación", "Lectura estratégica"], [
        ["agg_prevalencia_ubicacion", "agregados/", "Prevalencia por sede/territorio"],
        ["agg_prevalencia_edad", "agregados/", "Prevalencia por cohorte etaria"],
        ["agg_promedios_clinicos", "agregados/", "Control metabólico agregado"],
        ["agg_cohorte_riesgo", "agregados/", "Estratificación de riesgo"],
        ["agg_ocupacion_camas", "agregados/", "Capacidad / ocupación"],
        ["agg_productividad_medica", "negocio/", "Productividad por profesional"],
        ["agg_costo_servicio", "negocio/", "Costo vs facturado"],
        ["agg_tiempos_espera", "negocio/", "Espera promedio urgencias"],
        ["agg_medicamentos_top", "negocio/", "Ranking dispensación"],
        ["agg_margen_farmacia", "negocio/", "Margen por medicamento"],
    ])
    p(doc, "Apoyo: Dataset, Calidad diabetes, Predicción + Modelo ML. Criterio de cierre AGG: KPIs en UI + tablas agg_* trazables en MinIO.")

    h(doc, "6. ELT (puente de los tres niveles)")
    p(doc, "Sin ELT el operativo puede vivir; táctica compuesta y estratégico AGG quedan incompletos.")
    h(doc, "6.1 Extract / Load / Transform", 2)
    bullets(doc, [
        "Extract: generador/dataset, pacientes, hechos agenda/lab/farmacia/caja/urgencias/RRHH.",
        "Load: Parquet MinIO — stage/, negocio/ (+ agg_*), DWH hechos/dims/agregados/.",
        "Transform: limpieza; estrella; materialización AGG (prevalencia, margen, productividad, espera, rankings).",
    ])
    h(doc, "6.2 Relación con cada nivel", 2)
    table(doc, ["Nivel", "Consume", "Si ELT no corre"], [
        ["Operativo", "Parquet CRUD", "Sigue el día a día"],
        ["Táctico", "Workpanel + compuestos agg_*", "Compuestos vacíos"],
        ["Estratégico AGG", "KPIs + agg_*", "Franja AGG incompleta"],
    ])
    p(doc, "Orquestador: Datos → Orquestador (Airflow, DAG diabcare_elt). Demo: generar datos → pipeline → ver compuestos y AGG en Panel.")

    h(doc, "7. Trazabilidad nivel → pantalla")
    table(doc, ["Nivel", "Pantalla principal", "Evidencia"], [
        ["Operativo", "Módulos CRUD del menú", "Alta paciente, cita, atención, cobro"],
        ["Táctico", "Panel + PDF", "Workpanel, tarjetas, gráficas"],
        ["Estratégico AGG", "Panel franja AGG + Dataset/Calidad DM", "KPIs + agg_*"],
        ["ELT", "Orquestador + Dataset generador", "DAG / materialización"],
    ])

    h(doc, "8. Guía de demostración (orden sugerido)")
    bullets(doc, [
        "Arrancar backend (localhost:8000) y MinIO; login admin.",
        "Operativo (2–3 min): pacientes → agenda/mis citas → farmacia o caja.",
        "Táctico: Panel → Workpanel → simples → compuestos; Exportar PDF.",
        "Estratégico AGG: franja AGG; Dataset o Calidad diabetes si se quiere detalle.",
        "ELT: generador (si hace falta) → Orquestador Airflow.",
        "Cierre opcional: Predicción ML.",
    ])

    h(doc, "9. Repositorio y rutas")
    bullets(doc, [
        "Repo: https://github.com/Danxxx777/diabcare.git",
        "API: http://localhost:8000",
        "Panel: /paginas/clinico/analisis/informes/index.html",
        "PDF: /paginas/clinico/reportes/index.html",
        "Dataset: /paginas/datos/dataset/",
        "Orquestador: /paginas/datos/pipeline_elt/index.html",
        "Predicción: /paginas/clinico/prediccion/index.html",
        "Modelo ML: /paginas/datos/modelo_ml/index.html",
        "Calidad DM: /paginas/clinico/analisis/diabetes/index.html",
    ])

    h(doc, "10. Checklist de entrega")
    bullets(doc, [
        "Documento con Operativo + Táctico + Estratégico AGG + ELT (este archivo).",
        "Video en el mismo orden (link arriba); actualizar narración si el video viejo solo habla de Panel táctico.",
        "Sistema corriendo con datos: si AGG/compuestos en cero → regenerar dataset hospitalario.",
        "Capturas recomendadas: Panel (AGG+Workpanel+simples+compuestos), un CRUD, Orquestador.",
        "Narración clara: operativo=CRUD; táctico=Panel simples/compuestos; AGG=agregar/consolidados de dirección.",
    ])
    return doc

def main():
    doc = build()
    for out in OUTS:
        try:
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out)); print("OK", out)
        except Exception as e:
            print("FAIL", out, e)

if __name__ == "__main__":
    main()
