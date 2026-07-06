"""Corrige EVP 09.docx — punto 3 (casos de uso) sin referencias a entregas anteriores."""
from __future__ import annotations

from pathlib import Path

from docx import Document

SRC = Path(r"c:\Users\USER\Downloads\EVP 09.docx")
DST = Path(r"c:\Users\USER\Downloads\EVP 09-corregido.docx")


def replace_everywhere(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            if p.runs:
                p.runs[0].text = p.text.replace(old, new)
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = p.text.replace(old, new)
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        if p.runs:
                            p.runs[0].text = p.text.replace(old, new)
                            for r in p.runs[1:]:
                                r.text = ""
                        else:
                            p.text = p.text.replace(old, new)
                        n += 1
    return n


def replace_paragraph_exact(doc: Document, old: str, new: str, style: str | None = None) -> bool:
    for p in doc.paragraphs:
        if p.text.strip() == old:
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new
            if style:
                try:
                    p.style = doc.styles[style]
                except KeyError:
                    pass
            return True
    return False


def fix_matriz_permisos(doc: Document) -> None:
    rows_data = [
        ["Usuarios", "Si", "No", "No"],
        ["Pacientes / Consultas", "Si", "Si", "No"],
        ["Admisiones", "Si", "No", "No"],
        ["Agenda (citas)", "Si", "No", "No"],
        ["Mis citas", "No", "Si", "No"],
        ["Dataset/Pipeline", "Si", "No", "Si"],
        ["Predicción/Reportes", "Si", "Si", "Si"],
        ["Auditoria/Config", "Si", "No", "No"],
    ]
    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:4] != ["Modulo", "Admin", "Medico", "Analista"]:
            continue
        while len(table.rows) < len(rows_data) + 1:
            table.add_row()
        for i, data in enumerate(rows_data, 1):
            for j, val in enumerate(data):
                table.rows[i].cells[j].text = val
        break


def fix_cu_o02_tables(doc: Document) -> None:
    for table in doc.tables:
        flat = " | ".join(c.text for row in table.rows for c in row.cells)
        if "RF-O-P03-001" in flat and "RF-O-P03-002" in flat and "Filtrar por criterios" in flat:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 3 and cells[0] == "RF-O-P03-001":
                    row.cells[0].text = "RF-PAC-001"
                    row.cells[1].text = "Listar pacientes paginados"
                    row.cells[2].text = "GET /api/pacientes"
                if len(cells) >= 3 and cells[0] == "RF-O-P03-002":
                    row.cells[0].text = "RF-PAC-002"
                    row.cells[1].text = "Subir foto del paciente"
                    row.cells[2].text = "POST /api/pacientes/{id}/foto"
        if "CA-01" in flat and "CA-03" in flat and "Filtros aplicados" in flat:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0] == "CA-01":
                    row.cells[1].text = "Lista paginada de pacientes"
                if len(cells) >= 2 and cells[0] == "CA-02":
                    row.cells[1].text = "CRUD de paciente persiste en Parquet"
                if len(cells) >= 2 and cells[0] == "CA-03":
                    row.cells[1].text = "Foto visible cuando existe en MinIO"

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] == "Paquete" and "Registros P3" in cells[1]:
                row.cells[1].text = "Pacientes (HCE)"
            if len(cells) >= 3 and cells[0] == "RF-MC-001":
                row.cells[2].text = "GET /api/mis-citas"
            if len(cells) >= 3 and cells[0] == "RF-MC-002":
                row.cells[2].text = "PUT /api/mis-citas/{id}/estado"


def main() -> None:
    doc = Document(str(SRC))

    # APIs y typos
    api_fixes = [
        ("POST /api/agenda", "POST /api/citas"),
        ("GET /api/agenda", "GET /api/citas"),
        ("/api/agenda", "/api/citas"),
        ("GET /api/citas/mis-citas", "GET /api/mis-citas"),
        ("PUT /api/citas/{id}/estado", "PUT /api/mis-citas/{id}/estado"),
        ("data Waterhouse", "Data Warehouse"),
        ("Consulta externa → + Nueva consulta", "Consultas → + Nueva consulta"),
        ("1. Medico abre Consulta externa.", "1. El médico abre el módulo Consultas."),
    ]
    for old, new in api_fixes:
        replace_everywhere(doc, old, new)

    # Eliminar referencias a entregas / avances previos
    legacy_phrases = [
        ("definidos en el avance anterior", "del plan estratégico del proyecto"),
        ("objetivo definido en el avance anterior", "objetivo del proyecto"),
        ("meta definida en el avance anterior:", "meta del proyecto:"),
        ("avance estratégico anterior", "plan estratégico"),
        ("del avance estratégico anterior", "del plan estratégico"),
        ("avance anterior OO5.4.1", "OO5.4.1"),
        ("flujo avance anterior:", "flujo ELT:"),
        ("Capacidad analítica operativa OO5.5.1 del avance estratégico anterior.", "Capacidad analítica operativa OO5.5.1."),
        ("inteligencia clínica predictiva del avance estratégico anterior.", "inteligencia clínica predictiva (OO5.6.1)."),
        ("especificación del avance anterior (§13)", "especificación de alertas clínicas"),
        ("OO2.1.2 del avance estratégico anterior", "OO2.1.2"),
        (
            "El sistema DiabCare Hospital cumple la cadena de trazabilidad avance estratégico anterior → avance operativo actual con constitución v1.1.0, especificación operativa por paquetes, DWH de 63 tablas, once casos de uso demostrables (CU-O01–O10 y CU-O04-B) y suite de pruebas. Complete las figuras UML, capturas de pantalla y portada para la entrega final.",
            "El sistema DiabCare Hospital cumple la cadena de trazabilidad OE → OT → OO → paquetes → casos de uso, con constitución v1.1.0, especificación operativa, DWH de 63 tablas, casos de uso demostrables (CU-O01–O10 y CU-O04-B) y suite de pruebas. Complete las figuras UML, capturas de pantalla y portada para la entrega final.",
        ),
        (
            "El documento estratégico del avance anterior del proyecto establece la cadena BSC → OE → OT → OO y 59 casos de uso en tres niveles (CU-E, CU-T, CU-O). avance operativo actual consolida el nivel operativo implementado en el avance actual del sistema en ejecución, especificación SDD y UML.",
            "La documentación estratégica del proyecto establece la cadena BSC → OE → OT → OO y casos de uso en tres niveles (CU-E, CU-T, CU-O). Este entregable consolida el nivel operativo implementado: sistema en ejecución, especificación SDD y UML.",
        ),
        ("A.1 Flujo de valor del avance estratégico anterior", "A.1 Flujo de valor del sistema"),
        ("Objetivo operativo del avance anterior", "Objetivo operativo (OO)"),
        ("5.2 Checklist funcional del avance operativo", "5.2 Checklist funcional operativo"),
        ("plan estratégico → avance operativo actual", "plan estratégico → implementación operativa"),
        ("avance operativo actual", "implementación operativa"),
        ("RNF-002: Pipeline ELT 600K registros < 15 min (objetivo definido en el avance anterior).", "RNF-002: Pipeline ELT 600K registros < 15 min (objetivo del proyecto)."),
        ("RNF-003: Accuracy modelo RandomForest meta 96% (avance estratégico anterior).", "RNF-003: Accuracy modelo RandomForest meta 96%."),
        ("6. objetivo definido en el avance anterior: 600K registros < 15 minutos.", "6. Objetivo del proyecto: 600K registros < 15 minutos."),
        ("Demostración académica y pruebas de carga del DWH (avance anterior OO5.4.1).", "Demostración académica y pruebas de carga del DWH (OO5.4.1)."),
        ("Integridad del flujo avance anterior: PocketBase → Airflow → MinIO → API.", "Integridad del flujo ELT: PocketBase → Airflow → MinIO → API."),
        ("PocketBase solo origen (~100K); MinIO es BD; meta definida en el avance anterior: 600K registros < 15 min.", "PocketBase solo origen (~100K); MinIO es BD; meta del proyecto: 600K registros < 15 min."),
        ("Falta de trazabilidad entre objetivos estratégicos (definidos en el avance anterior) e implementación.", "Falta de trazabilidad entre objetivos estratégicos e implementación."),
    ]
    for old, new in legacy_phrases:
        replace_everywhere(doc, old, new)

    # Trazabilidad departamentos — separar usuarios de HCE
    replace_everywhere(doc, "CU-O01, CU-O02", "CU-O01, HU-O02 (P2)")
    replace_everywhere(doc, "Caso de uso: CU-O02. CRUD usuarios, asignación de roles", "Caso de uso: HU-O02 / P2. CRUD usuarios, asignación de roles")
    replace_everywhere(doc, "Caso de uso: CU-O02. Expediente, búsqueda, métricas por paciente", "Caso de uso: CU-O02. Expediente paciente (HCE), CRUD y foto")

    # CU-O02 flujo — reemplazos
    cu_o02_replacements = [
        ("Sesión JWT valida con permiso registros.", "Sesión JWT válida con permiso pacientes."),
        ("Existen registros clínicos en el DWH.", "Módulo Pacientes habilitado para el rol."),
        (
            "4. Se cargan encuentros, métricas (HbA1c, glucosa, IMC) y flags de riesgo.",
            "4. Consulta o edita datos demográficos del expediente (documento, sede, contacto).",
        ),
        (
            "5a. Opcional: sube o actualiza foto del paciente (almacenada en MinIO).",
            "5. Opcional: sube o actualiza foto del paciente (GET/POST /api/pacientes/{id}/foto, MinIO).",
        ),
        (
            "5b. El medico puede filtrar por diabetes, género o ubicación.",
            "6. El expediente queda disponible para admisiones, agenda y consultas clínicas.",
        ),
        (
            "6. Los datos provienen de hechos_diabetes y dimensiones asociadas.",
            "",
        ),
        (
            "CA-03\nExpediente muestra métricas del paciente",
            "CA-03\nFoto de paciente visible cuando existe en MinIO",
        ),
        (
            "RN-O-005: Alertas por umbrales HbA1c/glucosa.",
            "RN-PAC-001: Toda admisión o cita requiere paciente registrado previamente.",
        ),
    ]
    for old, new in cu_o02_replacements:
        if new:
            replace_everywhere(doc, old, new)
        else:
            # borrar párrafo exacto
            for p in doc.paragraphs:
                if p.text.strip() == old.strip():
                    p.text = ""

    # CU-O04 — quitar solapamiento no implementado
    cu_o04_replacements = [
        (
            "4. El sistema valida solapamiento de horarios.",
            "4. El sistema persiste la cita en estado programada.",
        ),
        (
            "Conflicto de horario → mensaje de error.",
            "Datos incompletos → HTTP 422 con detalle de campos.",
        ),
        ("CA-02\nValidación de conflicto funciona", "CA-02\nMédico asignado visible en la cita creada"),
        ("Catálogo de servicios y camas disponible.", "Catálogo de servicios y tipos de ingreso disponibles."),
    ]
    for old, new in cu_o04_replacements:
        replace_everywhere(doc, old, new)

    # CU-O03 precondición admisiones
    replace_everywhere(
        doc,
        "1. El medico accede al módulo Pacientes desde el menú Atención.",
        "1. El médico o administrador accede al módulo Pacientes desde Operaciones clínicas.",
    )
    replace_everywhere(
        doc,
        "3. El medico selecciona un paciente y abre el expediente.",
        "3. Selecciona un paciente y abre el expediente (o crea uno nuevo).",
    )

    fix_matriz_permisos(doc)
    fix_cu_o02_tables(doc)

    doc.save(str(DST))
    print(f"Guardado: {DST}")


if __name__ == "__main__":
    main()
