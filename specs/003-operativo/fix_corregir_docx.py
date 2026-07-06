"""Corrige corregir.docx y genera corregir-actualizado.docx con CU-O04-B completo."""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = Path(r"c:\Users\USER\Downloads\corregir.docx")
DST = Path(r"c:\Users\USER\Downloads\corregir-actualizado.docx")


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def replace_everywhere(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            if p.runs:
                p.runs[0].text = p.text.replace(old, new)
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = p.text.replace(old, new)
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        if p.runs:
                            p.runs[0].text = p.text.replace(old, new)
                            for r in p.runs[1:]:
                                r.text = ""
                        else:
                            p.text = p.text.replace(old, new)
                        n += 1
    return n


def fix_actors(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        paquete = None
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] == "Paquete":
                paquete = cells[1]
            if len(cells) >= 2 and cells[0] == "Actor principal":
                val = cells[1]
                if paquete == "Admisiones" and "administrador" in val.lower():
                    row.cells[1].text = "Administrador"
                    n += 1
                elif paquete == "Agenda" and "administrador" in val.lower():
                    row.cells[1].text = "Administrador"
                    n += 1
    return n


def fix_matriz_trazabilidad(doc: Document) -> None:
    mapping = {
        "CU-O02": "Pacientes HCE",
        "CU-O03": "Admisiones",
        "CU-O04": "Agenda / Citas",
        "CU-O05": "P3 Registros clínicos",
        "CU-O06": "P4 Dataset",
        "CU-O07": "P5 Análisis",
        "CU-O08": "P6 Predicción",
        "CU-O09": "P7 Reportes",
        "CU-O10": "P8 Pipeline ELT",
    }
    old_pkgs = {"P2", "P3", "P4", "P5", "P6", "P8", "P14"}
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells or cells[0] not in mapping:
                continue
            for cell in row.cells:
                if cell.text.strip() in old_pkgs:
                    cell.text = mapping[cells[0]]
                    break


def fix_matriz_permisos(doc: Document) -> None:
    rows_data = [
        ["Usuarios", "Si", "No", "No"],
        ["Pacientes / Consultas", "Si", "Si", "No"],
        ["Admisiones", "Si", "No", "No"],
        ["Agenda (citas)", "Si", "No", "No"],
        ["Mis citas", "No", "Si", "No"],
        ["Dataset/Pipeline", "Si", "No", "Si"],
        ["Predicción/Reportes", "Si", "Si", "Si"],
        ["Auditoria/Config", "Si", "No", "No"],
    ]
    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:4] != ["Modulo", "Admin", "Medico", "Analista"]:
            continue
        while len(table.rows) < len(rows_data) + 1:
            table.add_row()
        for i, data in enumerate(rows_data, 1):
            for j, val in enumerate(data):
                table.rows[i].cells[j].text = val
        break


def fix_cu_o02_o04_flows(doc: Document) -> None:
    replacements = {
        "CU-O02: Consultar expediente de paciente (HCE)": (
            "CU-O02: Gestionar expediente del paciente (HCE)"
        ),
        "1. El usuario abre el formulario de nueva admisión.": (
            "1. El administrador abre el formulario de nueva admisión."
        ),
        "2. Completa datos del paciente, servicio, motivo y prioridad.": (
            "2. Selecciona paciente, tipo de ingreso, servicio, médico tratante, motivo y prioridad."
        ),
        "1. Usuario accede al módulo Agenda.": (
            "1. El administrador accede al módulo Agenda."
        ),
        "3. Define paciente, medico, fecha, hora y tipo de consulta.": (
            "3. Selecciona paciente y médico asignado (listado de usuarios rol médico), fecha, hora y motivo."
        ),
        "Los siguientes diez casos de uso (CU-O01 a CU-O10) constituyen el núcleo operativo demostrable del sistema.": (
            "Los casos de uso CU-O01 a CU-O10 (más CU-O04-B Mis citas) constituyen el núcleo operativo demostrable del sistema."
        ),
        "diez casos de uso demostrables": "once casos de uso demostrables (CU-O01–O10 y CU-O04-B)",
        "CU-O02–O05, O07–O09": "CU-O02–O05, CU-O04-B, O07–O09",
        "Menú hospital + Admisiones + Agenda": "Menú hospital + Admisiones + Agenda + Mis citas (médico)",
    }
    for old, new in replacements.items():
        replace_everywhere(doc, old, new)

    # CU-O02: paso foto del paciente
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("4. Se cargan encuentros, métricas"):
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if "foto" not in nxt.lower():
                p.insert_paragraph_before(
                    "4. Opcional: sube o actualiza foto del paciente (GET/POST /api/pacientes/{id}/foto, MinIO).",
                    style="Normal",
                )
                break


def _style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles["Normal"]


def _make_table(doc: Document, template_table: Table, headers: list[str], rows: list[list[str]]) -> Table:
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    if template_table.style:
        table.style = template_table.style
    for ci, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[ci], h)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri].cells[ci], val)
    return table


def _para_element(text: str, style_name: str, doc: Document):
    p = doc.add_paragraph(text, style=style_name)
    el = p._element
    el.getparent().remove(el)
    return el


def _table_element(table: Table):
    el = table._tbl
    el.getparent().remove(el)
    return el


def insert_cu_o04b(doc: Document) -> bool:
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("CU-O05: Registrar consulta externa"):
            anchor = p
            break
    if anchor is None:
        return False

    # Evitar duplicar si ya existe
    for p in doc.paragraphs:
        if p.text.strip().startswith("CU-O04-B:"):
            return False

    tpl_attr = doc.tables[15]
    tpl_rf = doc.tables[16]
    tpl_ca = doc.tables[17]
    body = doc.element.body
    idx = body.index(anchor._element)

    blocks: list = []

    blocks.append(_para_element("CU-O04-B: Consultar y atender mis citas (médico)", "Heading 2", doc))

    t1 = _make_table(
        doc,
        tpl_attr,
        ["Atributo", "Valor"],
        [
            ["Paquete", "Citas / Mis citas"],
            ["Actor principal", "Médico"],
            ["OE / OT / OO", "OE4 / OT4.1 / OO5.2.1"],
            ["Ruta en aplicación", "Atención → Mis citas"],
        ],
    )
    blocks.append(_table_element(t1))

    for text, style in [
        ("Precondiciones", "Heading 3"),
        ("Sesión JWT válida con rol médico.", "List Bullet"),
        ("Existen citas asignadas por administración con el nombre del médico autenticado.", "List Bullet"),
        ("Módulo Mis citas habilitado en menú lateral.", "List Bullet"),
        ("Flujo principal", "Heading 3"),
        ("1. El médico accede al módulo Mis citas desde el menú Atención.", "Normal"),
        ("2. El sistema lista citas filtradas por su usuario (GET /api/citas/mis-citas).", "Normal"),
        ("3. El médico confirma una cita → estado confirmada.", "Normal"),
        ("4. Al pulsar Atender → estado atendida → redirección a Consultas (CU-O05).", "Normal"),
        ("5. Evento queda trazable en auditoría (P11).", "Normal"),
        ("Flujos alternos y excepciones", "Heading 3"),
        ("Paciente no acude → marcar estado no_asistio.", "List Bullet"),
        ("Cita no asignada al médico → HTTP 403.", "List Bullet"),
        ("Sin citas pendientes → mensaje informativo vacío.", "List Bullet"),
        ("Requisitos funcionales aplicables", "Heading 3"),
    ]:
        blocks.append(_para_element(text, style, doc))

    t2 = _make_table(
        doc,
        tpl_rf,
        ["ID", "Descripción", "Implementación"],
        [
            ["RF-MC-001", "Listar citas del médico autenticado", "GET /api/citas/mis-citas"],
            ["RF-MC-002", "Actualizar estado de cita", "PUT /api/citas/{id}/estado"],
        ],
    )
    blocks.append(_table_element(t2))
    blocks.append(_para_element("Reglas de negocio", "Heading 3", doc))
    blocks.append(_para_element(
        "RN-MC-001: El médico no crea ni cancela citas; solo la administración agenda (CU-O04).",
        "Normal",
        doc,
    ))
    blocks.append(_para_element(
        "RN-MC-002: Solo el médico asignado puede cambiar el estado de su cita.",
        "Normal",
        doc,
    ))
    blocks.append(_para_element("Criterios de aceptación", "Heading 3", doc))

    t3 = _make_table(
        doc,
        tpl_ca,
        ["ID", "Criterio"],
        [
            ["CA-01", "Listado muestra solo citas del médico autenticado"],
            ["CA-02", "Confirmar/atender actualiza estado correctamente"],
            ["CA-03", "Atender redirige a Consultas (CU-O05)"],
        ],
    )
    blocks.append(_table_element(t3))

    for el in reversed(blocks):
        body.insert(idx, el)

    return True


def fix_cu_o04_reglas(doc: Document) -> None:
    replace_everywhere(
        doc,
        "RN-AG-001: No permitir citas en horario ocupado.",
        "RN-AG-001: Solo administrador agenda citas; médico asignado desde listado rol médico.",
    )
    replace_everywhere(
        doc,
        "RN-ADM-002: Solo roles autorizados pueden admitir.",
        "RN-ADM-002: Solo administrador puede registrar admisiones.",
    )


def main() -> None:
    doc = Document(str(SRC))

    replace_everywhere(doc, "POST /api/agenda", "POST /api/citas")
    replace_everywhere(doc, "GET /api/agenda", "GET /api/citas")
    replace_everywhere(doc, "RF-AG-001", "RF-AG-001")
    replace_everywhere(doc, "Crear cita", "Crear cita")
    replace_everywhere(doc, "Consultar agenda", "Consultar citas")

    actors = fix_actors(doc)
    fix_matriz_trazabilidad(doc)
    fix_matriz_permisos(doc)
    fix_cu_o02_o04_flows(doc)
    fix_cu_o04_reglas(doc)
    inserted = insert_cu_o04b(doc)

    doc.save(str(DST))
    print(f"Guardado: {DST}")
    print(f"Actores corregidos: {actors}")
    print(f"CU-O04-B insertado: {inserted}")


if __name__ == "__main__":
    main()
