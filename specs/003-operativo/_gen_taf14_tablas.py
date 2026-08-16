# -*- coding: utf-8 -*-
"""Documento DiabCare: cobertura operativo / táctico / estratégico AGG.

Estructura propia del proyecto (no plantilla calcada del ejemplo ARME).
"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTS = [
    Path(r"D:\6to Software\Construcción de Software\TAF 14.docx"),
    Path(r"D:\6to Software\Construcción de Software\EVF 14.docx"),
    Path(r"C:\Users\USER\Downloads\TAF 14.docx"),
    Path(r"C:\Users\USER\Downloads\EVF 14.docx"),
]


def font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def para(doc, text, size=11, bold=False, after=8, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    font(r, size, bold)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        font(r, 13 if level == 1 else 11, True)
        r.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return h


def shade(cell, fill="E2E8F0"):
    tc = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    sh.set(qn("w:val"), "clear")
    tc.append(sh)


def cell(c, text, header=False, size=9):
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    font(r, size, header)
    if header:
        shade(c)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell(t.rows[0].cells[j], h, header=True)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            cell(t.rows[i].cells[j], str(v))
    doc.add_paragraph()


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    para(doc, "DiabCare Analytics", size=16, bold=True, center=True, after=4)
    para(
        doc,
        "Cobertura del sistema en niveles operativo, táctico y estratégico",
        size=12,
        bold=True,
        center=True,
        after=10,
    )
    para(doc, "Byron Loor Mendoza", size=11, after=2)
    para(doc, "Universidad Técnica Estatal de Quevedo — Carrera de Software", size=10, after=14)

    heading(doc, "1. Contexto")
    para(
        doc,
        (
            "DiabCare Analytics es la plataforma hospitalaria del proyecto para gestionar la "
            "atención de diabetes y leer su resultado clínico-administrativo. El trabajo diario "
            "queda en módulos de captura; la supervisión de periodo en el Panel; la mirada de "
            "dirección en indicadores ya agregados (AGG). Este documento describe esa cobertura "
            "con el vocabulario del propio sistema: módulos, Workpanel, informes simples, "
            "informes compuestos, franja Estratégico AGG, Dataset/Calidad diabetes, predicción "
            "y orquestación ELT."
        ),
        after=10,
    )

    heading(doc, "2. Tres formas de usar DiabCare")
    para(
        doc,
        (
            "No son tres aplicaciones distintas: son tres lecturas sobre la misma base en MinIO. "
            "Cada nivel responde a un tipo de pregunta y a un tipo de usuario."
        ),
        after=6,
    )
    table(
        doc,
        ["Nivel", "Pregunta central", "Dónde se vive en DiabCare", "Usuario típico"],
        [
            [
                "Operativo",
                "¿Qué registro atiendo ahora?",
                "Módulos CRUD (pacientes, agenda, lab, farmacia, caja, etc.)",
                "Recepcionista, médico, laboratorio, farmacia, caja",
            ],
            [
                "Táctico",
                "¿Cómo va el área en este periodo?",
                "Panel: Workpanel, informes simples y compuestos, exportación PDF",
                "Jefatura / coordinación",
            ],
            [
                "Estratégico (AGG)",
                "¿Qué dice el hospital ya consolidado?",
                "Franja Estratégico AGG, Dataset/Calidad DM, Predicción y Modelo ML",
                "Dirección / analítica",
            ],
        ],
    )
    para(
        doc,
        (
            "AGG significa agregar: totales, promedios, rankings y márgenes. No es un sinónimo "
            "de “pantalla bonita”; es el resultado de materializar hechos operativos en tablas "
            "y KPIs consolidados."
        ),
        after=10,
    )

    heading(doc, "3. Nivel operativo")
    para(
        doc,
        (
            "El operativo cubre el ciclo asistencial y administrativo del día. Cada módulo "
            "escribe o consulta registros puntuales. Aquí no se entrena un modelo: la decisión "
            "es transaccional (guardar, dispensar, cobrar, admitir, confirmar)."
        ),
        after=6,
    )
    table(
        doc,
        ["Módulo", "Qué permite hacer", "Decisión típica"],
        [
            ["Autenticación", "Iniciar y cerrar sesión con rol", "Permitir o denegar el acceso"],
            ["Pacientes", "Mantener HCE y foto del paciente", "Alta, corrección o búsqueda del expediente"],
            ["Admisiones", "Registrar ingresos y egresos", "Admitir, dar alta o cancelar"],
            ["Agenda / citas", "Programar y controlar turnos (admin)", "Agendar, reprogramar, cancelar o cobrar"],
            ["Mis citas", "Bandeja del médico", "Confirmar cita o pasar a atender"],
            ["Registros clínicos", "Documentar consulta y métricas", "Guardar HbA1c, glucosa, notas, sede"],
            ["Urgencias", "Atender demanda no programada", "Triar, atender o cerrar el caso"],
            ["Laboratorio", "Órdenes y resultados", "Solicitar, cargar resultado o completar orden"],
            ["Comorbilidades", "Registrar complicaciones asociadas", "Alta o edición por paciente"],
            ["Recetas", "Prescribir tratamiento", "Emitir o anular receta"],
            ["Farmacia", "Dispensar e inventario", "Entregar medicamento o ajustar stock"],
            ["Facturación / caja", "Facturar y cobrar", "Emitir, cobrar, anular o imprimir recibo"],
            ["Usuarios", "Gobernar cuentas y roles", "Crear, aprobar, cambiar rol o desactivar"],
            ["Auditoría", "Consultar bitácora de acciones", "Rastrear quién hizo qué y cuándo"],
            ["Notificaciones", "Avisos operativos/clínicos", "Leer y marcar pendientes"],
        ],
    )

    heading(doc, "4. Nivel táctico")
    para(
        doc,
        (
            "El Panel concentra la supervisión de periodo. El Workpanel da el pulso general; "
            "los informes simples resumen un área; los compuestos ya muestran series y tablas "
            "agregadas (ingresos por día, top medicamentos, productividad, margen, costo de "
            "servicio, tiempos de espera). La jefatura decide con dashboard; no con IA."
        ),
        after=6,
    )

    heading(doc, "4.1 Workpanel e informes simples", level=2)
    table(
        doc,
        ["Vista", "Qué muestra", "Para qué sirve"],
        [
            ["Workpanel", "Total pacientes, con/sin DM, prevalencia", "Elegir por dónde empezar la revisión"],
            ["Agenda / Mis citas", "Turnos, estados y cobro del día", "Empujar agenda o cobros atrasados"],
            ["Urgencias", "Casos, espera y atendidos", "Priorizar triage o personal"],
            ["Laboratorio", "Pendientes, completadas, resultados", "Acelerar órdenes atrasadas"],
            ["Facturación / caja", "Vigentes, pagadas, monto pendiente", "Conciliar o perseguir cobros"],
            ["Pacientes", "Activos / inactivos del padrón", "Depurar o completar expedientes"],
            ["Admisiones", "Hospitalizados, altas, cancelados", "Revisar ocupación y egresos"],
            ["Farmacia", "Recetas, stock bajo, ventas", "Reponer o corregir inventario"],
            ["RRHH", "Personal e ingreso asociado", "Revisar productividad de personal"],
            ["Comorbilidades", "Complicaciones registradas", "Ver carga de complicaciones del periodo"],
            ["Exportar PDF", "Informe simple, compuesto o completo", "Archivar o enviar fuera del sistema"],
        ],
    )

    heading(doc, "4.2 Informes compuestos", level=2)
    table(
        doc,
        ["Informe", "Indicador / tabla", "Lectura de jefatura"],
        [
            ["Ingresos por día", "Serie de recaudación", "Detectar caídas o picos del periodo"],
            ["Top medicamentos", "agg_medicamentos_top", "Priorizar compras de alto movimiento"],
            ["Productividad médica", "agg_productividad_medica", "Rebalancear agenda entre médicos"],
            ["Margen por medicamento", "agg_margen_farmacia", "Ajustar precio o costo de ítems flojos"],
            ["Costo de servicio", "agg_costo_servicio", "Corregir tarifas o costos deficitarios"],
            ["Tiempos de espera", "agg_tiempos_espera", "Ampliar capacidad si la demora sube"],
        ],
    )

    heading(doc, "5. Nivel estratégico (AGG)")
    para(
        doc,
        (
            "Aquí ya no se decide sobre una fila: se lee el hospital consolidado. La franja "
            "Estratégico AGG del Panel, las tablas agg_* y Calidad diabetes responden con "
            "totales, promedios y tasas. Dirección usa estos números para metas y prioridades "
            "del siguiente periodo."
        ),
        after=6,
    )
    table(
        doc,
        ["Indicador agregado", "Fuente en DiabCare", "Decisión de dirección"],
        [
            ["Facturado total", "KPI AGG del Panel", "Fijar o revisar meta de recaudación"],
            ["Margen de farmacia", "KPI AGG + agg_margen_farmacia", "Redefinir compras o precios a escala"],
            ["Espera promedio (urgencias)", "KPI AGG + agg_tiempos_espera", "Invertir capacidad si el agregado empeora"],
            ["Productividad (consultas)", "KPI AGG de productividad", "Planear agenda o contratación"],
            ["Laboratorio pendiente", "KPI AGG de lab", "Reforzar diagnóstico a nivel hospital"],
            ["Prevalencia / cohorte DM", "Calidad diabetes · Dataset DWH", "Priorizar sedes o grupos etarios"],
            ["Control metabólico (HbA1c / glucosa)", "KPIs y gráficos de Calidad diabetes", "Definir campañas o protocolos de control"],
            ["Frescura del dato analítico", "Orquestador (Airflow) · Dataset", "Reprocesar ELT o regenerar stage/DWH"],
        ],
    )

    heading(doc, "6. Apoyo de inteligencia artificial")
    para(
        doc,
        (
            "La IA no sustituye al mostrador ni al Panel. En DiabCare aparece solo cuando hay "
            "que estimar un riesgo o gobernar el modelo, siempre sobre variables clínicas ya "
            "disponibles tras el flujo de datos."
        ),
        after=6,
    )
    table(
        doc,
        ["Pantalla", "Qué responde", "Modelo", "Uso"],
        [
            [
                "Predicción",
                "Probabilidad y nivel de riesgo del paciente",
                "Random Forest",
                "Focalizar seguimiento clínico-educativo",
            ],
            [
                "Modelo ML",
                "Si el modelo sigue siendo usable (métricas / estado)",
                "Random Forest (evaluación y reentrenamiento)",
                "Reentrenar, congelar o seguir usando el modelo",
            ],
        ],
    )

    heading(doc, "7. Papel del ELT")
    para(
        doc,
        (
            "El ELT (extraer → cargar → transformar hacia MinIO/Parquet) conecta lo escrito en "
            "operativo con lo que se lee en táctico y en AGG. Sin pipeline fresco, el Panel y "
            "las tablas agregadas se quedan atrás del día a día. Por eso el Orquestador "
            "(Airflow) y el Dataset forman parte de la cobertura estratégica: no son un cuarto "
            "nivel aparte, son la condición para que AGG tenga sentido."
        ),
        after=10,
    )

    heading(doc, "8. Conclusión")
    para(
        doc,
        (
            "DiabCare cubre los tres niveles con piezas distintas pero encadenadas: módulos "
            "para ejecutar, Panel para supervisar el periodo, AGG e IA para la mirada "
            "institucional y predictiva. El valor del diseño está en esa separación: cada "
            "usuario ve la pregunta que le corresponde, sobre la misma base de datos."
        ),
        after=10,
    )

    para(doc, "Referencias", size=11, bold=True, after=4)
    para(doc, "Repositorio: https://github.com/Danxxx777/diabcare.git", size=10, after=2)
    para(doc, "Video: https://youtu.be/3425-TO3UN4", size=10, after=2)
    return doc


def main():
    doc = build()
    for out in OUTS:
        try:
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out))
            print("OK", out)
        except Exception as e:
            print("FAIL", out, e)


if __name__ == "__main__":
    main()
